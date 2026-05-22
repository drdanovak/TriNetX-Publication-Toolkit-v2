"""
14_📄_Manuscript_Assembly.py — Phase 4 capstone tool.

Pulls the study context, registered TriNetX uploads, and Pre-Analysis Plan
answers from the current session into a single Word document containing
draft Methods and Results paragraphs ready to paste into a manuscript.

The author fills in three or four prose blocks (population, exposure,
comparator, primary outcome). Everything else — cohort sizes, event counts,
risks, RRs, HRs, p-values, PH assessments — is parsed from the TriNetX
files and dropped into the text in correct journal style.
"""

from __future__ import annotations

import io
from datetime import datetime
from io import BytesIO
from typing import Dict, List, Optional, Tuple

import pandas as pd
import streamlit as st

from utils.session import (
    ensure_context,
    render_context_banner,
    render_methods_text,
    render_check_callout,
    list_uploads,
)
from utils.parsers import (
    parse_moa_csv, parse_km_csv,
    EXPORT_BASELINE, EXPORT_MOA, EXPORT_KM,
    MOAExport, KMExport,
)
from utils.formatters import (
    fmt_p, fmt_pct, fmt_ratio_with_ci, fmt_count, crosses_null,
)
from utils.exports import render_reproducibility_footer, _stamp_filename, DOCX_AVAILABLE

if DOCX_AVAILABLE:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

st.set_page_config(page_title="Manuscript Assembly", page_icon="📄", layout="wide")
ctx = ensure_context()

st.title("Manuscript Assembly")
st.caption(
    "Generates a Word document with Methods and Results paragraphs ready to paste into your "
    "manuscript. Pulls parsed values from every TriNetX file uploaded this session."
)

render_context_banner()

if not DOCX_AVAILABLE:
    st.error("python-docx is not installed. Run `pip install python-docx` and reload.")
    st.stop()

# ---------------------------------------------------------------------------
# Pull session state
# ---------------------------------------------------------------------------

moa_uploads = list_uploads(of_type=EXPORT_MOA)
km_uploads = list_uploads(of_type=EXPORT_KM)
baseline_uploads = list_uploads(of_type=EXPORT_BASELINE)
pap = st.session_state.get("pap_v2", {}) or {}

if not (moa_uploads or km_uploads):
    st.info(
        "Upload at least one MOA or KM file in the Outcomes Table, Forest Plot, or KM "
        "Diagnostics tool first. Once a file is uploaded anywhere in the toolkit, this "
        "assembly tool can use it."
    )
    st.stop()

# Re-parse
moas: List[Tuple[str, MOAExport]] = []
kms: List[Tuple[str, KMExport]] = []
for u in moa_uploads:
    try:
        moas.append((u["name"], parse_moa_csv(BytesIO(u["bytes"]))))
    except Exception:
        pass
for u in km_uploads:
    try:
        kms.append((u["name"], parse_km_csv(BytesIO(u["bytes"]))))
    except Exception:
        pass

st.caption(
    f"Using {len(moas)} MOA, {len(kms)} KM, and {len(baseline_uploads)} Baseline file(s) "
    f"from this session."
    + (" Pre-Analysis Plan answers detected." if pap else "")
)

# ---------------------------------------------------------------------------
# Author fills in the three or four prose blocks the toolkit cannot infer
# ---------------------------------------------------------------------------

st.markdown("### Fill in what the toolkit cannot infer")
st.caption("Everything else comes from your uploaded files and study context.")

col1, col2 = st.columns(2)
with col1:
    population_text = st.text_area(
        "Study population (one or two sentences)",
        value=pap.get("population", ""),
        height=80,
        placeholder="Adults aged 18+ with a recorded diagnosis of...",
    )
    exposure_text = st.text_area(
        "Exposure / intervention",
        value=pap.get("intervention", ""),
        height=80,
    )
with col2:
    comparator_text = st.text_area(
        "Comparator",
        value=pap.get("comparator", ""),
        height=80,
    )
    primary_outcome_text = st.text_area(
        "Primary outcome",
        value=pap.get("outcome_primary", ""),
        height=80,
    )

manuscript_title = st.text_input("Manuscript title (optional)", value=ctx.study_title or "")

# ---------------------------------------------------------------------------
# Build the prose
# ---------------------------------------------------------------------------

def _ph(value: str, fallback: str) -> str:
    return value.strip() if value and value.strip() else f"[{fallback}]"


def _fmt_p_inline(p) -> str:
    s = fmt_p(p, style=ctx.p_style, na="")
    if not s:
        return "not reported"
    return f"p {s}" if s.startswith("<") else f"p = {s}"


def build_methods() -> List[Dict[str, str]]:
    paras = []

    # Data source
    paras.append({
        "heading": "Data source",
        "text": (
            "We conducted a retrospective observational study using de-identified electronic "
            "health record data from the TriNetX research network. Analyses were performed in "
            "TriNetX Analytics, and manuscript-ready tables, figures, and rigor diagnostics "
            "were prepared with the TriNetX Publication Toolkit (v2)."
        ),
    })

    # Population
    paras.append({
        "heading": "Study population",
        "text": (
            f"{_ph(population_text, 'Description of the study population')} The "
            f"{ctx.treated_label} cohort (treated/exposed) was compared with the "
            f"{ctx.control_label} cohort. Exposure was defined as: "
            f"{_ph(exposure_text, 'exposure definition')}. The comparator was defined as: "
            f"{_ph(comparator_text, 'comparator definition')}."
        ),
    })

    # Outcomes
    outcome_para = f"The primary outcome was {_ph(primary_outcome_text, 'primary outcome')}."
    if pap.get("outcome_secondary"):
        outcome_para += f" Secondary outcomes: {pap['outcome_secondary']}"
    paras.append({"heading": "Outcomes", "text": outcome_para})

    # Time windows (if PAP completed)
    if pap.get("index_definition") or pap.get("lookback_window") or pap.get("followup_window"):
        windows = (
            f"The index date was {_ph(str(pap.get('index_definition', '')), 'index date definition')}. "
            f"Baseline covariates were ascertained over a {pap.get('lookback_window', '[look-back]')}-day "
            f"look-back, with a {pap.get('washout_window', '[washout]')}-day washout to enforce a "
            f"new-user design. Patients were followed for up to {pap.get('followup_window', '[follow-up]')} days."
        )
        paras.append({"heading": "Time windows", "text": windows})

    # Statistical analysis
    stat = (
        f"The significance threshold was α = {ctx.alpha} ({'two-sided' if ctx.two_sided else 'one-sided'}). "
        f"Effect estimates are reported with 95% confidence intervals."
    )
    if pap.get("matching_strategy"):
        stat = f"Cohorts were balanced using {pap['matching_strategy'].lower()}. " + stat
    if pap.get("multiple_comparison_method") and "None" not in pap["multiple_comparison_method"]:
        stat += f" Multiple-comparisons correction: {pap['multiple_comparison_method']}."
    paras.append({"heading": "Statistical analysis", "text": stat})

    # Reporting
    paras.append({
        "heading": "Reporting",
        "text": (
            "Reporting follows the STROBE guidelines for observational studies and the RECORD "
            "extension for studies using routinely-collected health data."
        ),
    })

    return paras


def build_results() -> List[Dict[str, str]]:
    paras = []
    treated_first = ctx.treated_cohort == "1"

    # Cohort sizes (from first MOA file if available)
    if moas:
        _, first = moas[0]
        paras.append({
            "heading": "Cohort sizes",
            "text": (
                f"After cohort construction and matching, the {ctx.treated_label} cohort comprised "
                f"{fmt_count(first.cohort_1_n if treated_first else first.cohort_2_n)} patients and "
                f"the {ctx.control_label} cohort comprised "
                f"{fmt_count(first.cohort_2_n if treated_first else first.cohort_1_n)} patients."
            ),
        })

    # One paragraph per binary outcome
    for name, m in moas:
        outcome_label = name.rsplit(".", 1)[0].replace("_", " ").strip()
        t_risk = m.cohort_1_risk if treated_first else m.cohort_2_risk
        c_risk = m.cohort_2_risk if treated_first else m.cohort_1_risk
        t_events = m.cohort_1_events if treated_first else m.cohort_2_events
        c_events = m.cohort_2_events if treated_first else m.cohort_1_events

        rr = fmt_ratio_with_ci(m.risk_ratio_value, m.risk_ratio_ci[0], m.risk_ratio_ci[1],
                                decimals=ctx.ratio_decimals, na="not estimable")
        p_str = _fmt_p_inline(m.primary_p_value)

        if m.risk_ratio_value is not None and m.risk_ratio_value < 1.0:
            direction = "lower"
        elif m.risk_ratio_value is not None and m.risk_ratio_value > 1.0:
            direction = "higher"
        else:
            direction = "comparable"

        sig = ""
        if m.risk_ratio_ci[0] is not None and m.risk_ratio_ci[1] is not None:
            sig = (" The 95% confidence interval did not cross the null."
                   if not crosses_null(m.risk_ratio_ci[0], m.risk_ratio_ci[1])
                   else " The 95% confidence interval crossed the null.")

        text = (
            f"For {outcome_label}, the risk in the {ctx.treated_label} cohort was "
            f"{fmt_pct(t_risk, decimals=ctx.pct_decimals)} ({fmt_count(t_events)} events) versus "
            f"{fmt_pct(c_risk, decimals=ctx.pct_decimals)} ({fmt_count(c_events)} events) in the "
            f"{ctx.control_label} cohort. The risk was {direction} in the {ctx.treated_label} "
            f"cohort (risk ratio {rr}; {p_str}).{sig}"
        )
        paras.append({"heading": outcome_label, "text": text})

    # One paragraph per time-to-event outcome
    for name, k in kms:
        outcome_label = name.rsplit(".", 1)[0].replace("_", " ").strip()
        hr = fmt_ratio_with_ci(k.hazard_ratio, k.hazard_ratio_ci[0], k.hazard_ratio_ci[1],
                                decimals=ctx.ratio_decimals, na="not estimable")
        lr_p = _fmt_p_inline(k.log_rank_p)

        ph_clause = ""
        if k.ph_assumption_p is not None:
            ph_p = _fmt_p_inline(k.ph_assumption_p)
            if k.ph_assumption_p < 0.05:
                ph_clause = (
                    f" The proportional hazards assumption was not supported ({ph_p}); the hazard "
                    f"ratio is interpreted as a time-averaged estimate."
                )
            else:
                ph_clause = f" The proportional hazards assumption was supported ({ph_p})."

        paras.append({
            "heading": f"Time-to-event: {outcome_label}",
            "text": f"The hazard ratio for {outcome_label} was {hr} (log-rank {lr_p}).{ph_clause}",
        })

    if not paras:
        paras.append({
            "heading": "Results",
            "text": "[No outcomes parsed. Upload MOA or KM files in another tool first.]",
        })

    return paras


methods = build_methods()
results = build_results()

# ---------------------------------------------------------------------------
# Preview
# ---------------------------------------------------------------------------

st.markdown("### Preview")

st.markdown("#### Methods")
for p in methods:
    st.markdown(f"**{p['heading']}.** {p['text']}")
    st.markdown("")

st.markdown("#### Results")
for p in results:
    st.markdown(f"**{p['heading']}.** {p['text']}")
    st.markdown("")

# ---------------------------------------------------------------------------
# Build and offer the DOCX
# ---------------------------------------------------------------------------

def build_docx() -> bytes:
    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Inches(1)

    if manuscript_title.strip():
        t = doc.add_paragraph()
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = t.add_run(manuscript_title.strip())
        r.bold = True
        r.font.name = "Arial"
        r.font.size = Pt(14)
        doc.add_paragraph()

    # Placeholder note
    note = doc.add_paragraph()
    nr = note.add_run(
        "Bracketed text marks placeholders the toolkit could not infer. Replace before submission."
    )
    nr.italic = True
    nr.font.name = "Times New Roman"
    nr.font.size = Pt(9)
    nr.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
    doc.add_paragraph()

    def add_body(text: str, bold_lead: Optional[str] = None):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 2.0
        if bold_lead:
            r = p.add_run(bold_lead + ". ")
            r.bold = True
            r.font.name = "Times New Roman"
            r.font.size = Pt(11)
        body = p.add_run(text)
        body.font.name = "Times New Roman"
        body.font.size = Pt(11)

    def add_heading(text: str):
        h = doc.add_heading(text, level=1)
        for r in h.runs:
            r.font.name = "Arial"
            r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    add_heading("Methods")
    for p in methods:
        add_body(p["text"], bold_lead=p["heading"])

    add_heading("Results")
    for p in results:
        add_body(p["text"], bold_lead=p["heading"])

    # Brief footer with source list
    add_heading("Source files")
    add_body(
        "The following TriNetX exports underlie the values reported above: "
        + "; ".join(
            [n for n, _ in moas] + [n for n, _ in kms] + [u["name"] for u in baseline_uploads]
        )
        + ".",
    )

    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = foot.add_run(
        f"Draft assembled by the TriNetX Publication Toolkit v2 on "
        f"{datetime.now().strftime('%Y-%m-%d %H:%M')}."
    )
    fr.italic = True
    fr.font.name = "Times New Roman"
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


st.markdown("### Download")

st.download_button(
    "Download manuscript draft (DOCX)",
    data=build_docx(),
    file_name=_stamp_filename("manuscript_draft", "docx"),
    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    use_container_width=True,
    type="primary",
)

# ---------------------------------------------------------------------------
# Footer trio
# ---------------------------------------------------------------------------

render_check_callout(
    "This is a draft, not a finished manuscript. Verify that every parsed value in the Results "
    "section matches the corresponding TriNetX export, and replace every bracketed placeholder "
    "before submission."
)

render_methods_text(
    """A manuscript draft was assembled by the Manuscript Assembly tool of the TriNetX
Publication Toolkit (v2). The tool re-parses TriNetX exports uploaded during the session
and inserts cohort sizes, event counts, risks, effect estimates with 95% confidence
intervals, and p-values directly into Methods and Results paragraphs.""",
    title="Methods text describing the assembly itself",
)

render_reproducibility_footer("Manuscript Assembly")
