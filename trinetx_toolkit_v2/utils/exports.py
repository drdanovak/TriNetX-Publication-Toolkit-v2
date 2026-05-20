"""
exports.py — Shared export helpers for the TriNetX Publication Toolkit.

Each tool that produces a figure should offer PNG, SVG, and (where applicable)
DOCX/CSV downloads using the same UI affordances. The `render_figure_downloads`
helper places three or four download buttons in a row with consistent labels
and file naming.
"""

from __future__ import annotations

import io
import re
from datetime import datetime
from typing import Optional

import matplotlib.pyplot as plt
import pandas as pd
import streamlit as st

from .figure_defaults import DEFAULT_DPI

try:
    from docx import Document
    from docx.shared import Pt, Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def _slugify(name: str) -> str:
    name = re.sub(r"[^A-Za-z0-9_-]+", "_", name).strip("_")
    return name or "trinetx_output"


def _stamp_filename(prefix: str, ext: str) -> str:
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    return f"{_slugify(prefix)}_{ts}.{ext}"


def figure_to_png_bytes(fig: plt.Figure, dpi: int = DEFAULT_DPI) -> bytes:
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=dpi, bbox_inches="tight")
    buf.seek(0)
    return buf.getvalue()


def figure_to_svg_bytes(fig: plt.Figure) -> bytes:
    buf = io.BytesIO()
    fig.savefig(buf, format="svg", bbox_inches="tight")
    buf.seek(0)
    return buf.getvalue()


def render_figure_downloads(
    fig: plt.Figure,
    prefix: str,
    dpi: int = DEFAULT_DPI,
    include_pdf: bool = False,
) -> None:
    """Render standardized PNG / SVG (and optional PDF) download buttons."""
    cols = st.columns(3 if include_pdf else 2)

    with cols[0]:
        st.download_button(
            label="Download PNG",
            data=figure_to_png_bytes(fig, dpi=dpi),
            file_name=_stamp_filename(prefix, "png"),
            mime="image/png",
            use_container_width=True,
        )
    with cols[1]:
        st.download_button(
            label="Download SVG",
            data=figure_to_svg_bytes(fig),
            file_name=_stamp_filename(prefix, "svg"),
            mime="image/svg+xml",
            use_container_width=True,
        )
    if include_pdf:
        with cols[2]:
            buf = io.BytesIO()
            fig.savefig(buf, format="pdf", bbox_inches="tight")
            buf.seek(0)
            st.download_button(
                label="Download PDF",
                data=buf.getvalue(),
                file_name=_stamp_filename(prefix, "pdf"),
                mime="application/pdf",
                use_container_width=True,
            )


def render_table_downloads(
    df: pd.DataFrame,
    prefix: str,
    docx_title: Optional[str] = None,
    include_docx: bool = True,
) -> None:
    """Render CSV (and optional DOCX) download buttons for a tabular output."""
    cols = st.columns(2 if include_docx and DOCX_AVAILABLE else 1)

    with cols[0]:
        st.download_button(
            label="Download CSV",
            data=df.to_csv(index=False).encode("utf-8"),
            file_name=_stamp_filename(prefix, "csv"),
            mime="text/csv",
            use_container_width=True,
        )

    if include_docx and DOCX_AVAILABLE and len(cols) > 1:
        with cols[1]:
            st.download_button(
                label="Download DOCX",
                data=dataframe_to_docx_bytes(df, title=docx_title or prefix),
                file_name=_stamp_filename(prefix, "docx"),
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )


def dataframe_to_docx_bytes(df: pd.DataFrame, title: Optional[str] = None) -> bytes:
    """Render a DataFrame as a clean DOCX table."""
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx is not available in this environment.")
    doc = Document()
    if title:
        h = doc.add_heading(title, level=1)
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr[i].text = str(col)
    for _, row in df.iterrows():
        rc = table.add_row().cells
        for i, col in enumerate(df.columns):
            rc[i].text = "" if pd.isna(row[col]) else str(row[col])
    for r in table.rows:
        for c in r.cells:
            for p in c.paragraphs:
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(10)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def render_reproducibility_footer(tool_name: str, version: str = "2.0.0") -> None:
    """Print a small footer with toolkit version and timestamp."""
    ts = datetime.now().strftime("%Y-%m-%d %H:%M")
    st.caption(
        f"TriNetX Publication Toolkit v{version} &middot; {tool_name} &middot; "
        f"Generated {ts}"
    )
